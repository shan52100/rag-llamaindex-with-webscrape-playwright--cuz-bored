import os
import io
import json
import asyncio
from typing import List, Dict, Any
from concurrent.futures import ThreadPoolExecutor
from dotenv import load_dotenv

load_dotenv()

from llama_index.core import Document, VectorStoreIndex, Settings, StorageContext
from llama_index.core.node_parser import SentenceSplitter
from llama_index.llms.groq import Groq
from llama_index.embeddings.fastembed import FastEmbedEmbedding
from llama_index.vector_stores.qdrant import QdrantVectorStore
from qdrant_client import QdrantClient
from qdrant_client.models import Distance, VectorParams

EMBED_DIM = 384  # BAAI/bge-small-en-v1.5 output dimension


class RAGEngine:
    def __init__(self):
        self.llm = Groq(
            model="llama-3.3-70b-versatile",
            api_key=os.getenv("GROQ_API_KEY"),
        )
        self.embed_model = FastEmbedEmbedding(model_name="BAAI/bge-small-en-v1.5")

        Settings.llm = self.llm
        Settings.embed_model = self.embed_model
        Settings.chunk_size = 512
        Settings.chunk_overlap = 50

        self.qdrant = QdrantClient(
            url=os.getenv("QDRANT_URL"),
            api_key=os.getenv("QDRANT_API_KEY"),
        )

        self.splitter = SentenceSplitter(chunk_size=512, chunk_overlap=50)
        self._executor = ThreadPoolExecutor(max_workers=3)
        self._sources: Dict[str, List[str]] = {}

    # ── helpers ──────────────────────────────────────────────────────────────

    def _ensure_collection(self, name: str):
        existing = {c.name for c in self.qdrant.get_collections().collections}
        if name not in existing:
            self.qdrant.create_collection(
                collection_name=name,
                vectors_config=VectorParams(size=EMBED_DIM, distance=Distance.COSINE),
            )

    def _parse_file(self, content: bytes, filename: str) -> str:
        ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else "txt"
        if ext == "pdf":
            try:
                from pypdf import PdfReader
                reader = PdfReader(io.BytesIO(content))
                return "\n".join(page.extract_text() or "" for page in reader.pages)
            except Exception:
                pass
        elif ext == "docx":
            try:
                import docx
                doc = docx.Document(io.BytesIO(content))
                return "\n".join(p.text for p in doc.paragraphs)
            except Exception:
                pass
        return content.decode("utf-8", errors="ignore")

    # ── sync internals (run in thread executor) ───────────────────────────────

    def _sync_ingest(self, text: str, source: str, collection_name: str):
        self._ensure_collection(collection_name)
        doc = Document(text=text, metadata={"source": source})
        nodes = self.splitter.get_nodes_from_documents([doc])

        vector_store = QdrantVectorStore(
            client=self.qdrant, collection_name=collection_name
        )
        storage_ctx = StorageContext.from_defaults(vector_store=vector_store)
        VectorStoreIndex(nodes, storage_context=storage_ctx, show_progress=False)

        self._sources.setdefault(collection_name, [])
        if source not in self._sources[collection_name]:
            self._sources[collection_name].append(source)

        return nodes

    def _sync_query(self, message: str, collection_name: str) -> str:
        existing = {c.name for c in self.qdrant.get_collections().collections}
        if collection_name not in existing:
            return (
                f"No collection named '{collection_name}' exists yet. "
                "Please go to Upload / Scrape, add a document or URL, "
                "then select that collection from the Source dropdown before chatting."
            )
        vector_store = QdrantVectorStore(
            client=self.qdrant, collection_name=collection_name
        )
        index = VectorStoreIndex.from_vector_store(
            vector_store=vector_store, embed_model=self.embed_model
        )
        engine = index.as_query_engine(llm=self.llm, similarity_top_k=5)
        return str(engine.query(message))

    # ── public async API ──────────────────────────────────────────────────────

    async def ingest_text(self, text: str, source: str, collection_name: str) -> Dict:
        loop = asyncio.get_running_loop()
        nodes = await loop.run_in_executor(
            self._executor, self._sync_ingest, text, source, collection_name
        )
        return {
            "collection": collection_name,
            "source": source,
            "total_chunks": len(nodes),
            "chunks": [
                {
                    "index": i,
                    "text": node.text,
                    "char_count": len(node.text),
                    "metadata": node.metadata,
                }
                for i, node in enumerate(nodes)
            ],
        }

    async def ingest_document(
        self, content: bytes, filename: str, collection_name: str
    ) -> Dict:
        text = self._parse_file(content, filename)
        return await self.ingest_text(text, filename, collection_name)

    async def query(self, message: str, collection_name: str = "default") -> str:
        loop = asyncio.get_running_loop()
        return await loop.run_in_executor(
            self._executor, self._sync_query, message, collection_name
        )

    async def get_chunks(self, collection_name: str) -> List[Dict]:
        try:
            result = self.qdrant.scroll(
                collection_name=collection_name,
                limit=200,
                with_payload=True,
                with_vectors=False,
            )
            chunks = []
            for i, point in enumerate(result[0]):
                payload = point.payload or {}
                raw = payload.get("_node_content", "{}")
                try:
                    node_data = json.loads(raw) if isinstance(raw, str) else raw
                    text = node_data.get("text", "")
                    metadata = node_data.get("metadata", {})
                except Exception:
                    text = str(raw)
                    metadata = {}
                chunks.append(
                    {
                        "id": str(point.id),
                        "index": i,
                        "text": text,
                        "char_count": len(text),
                        "metadata": metadata,
                    }
                )
            return chunks
        except Exception:
            return []

    async def list_sources(self) -> List[Dict]:
        try:
            collections = self.qdrant.get_collections().collections
            out = []
            for c in collections:
                info = self.qdrant.get_collection(c.name)
                out.append(
                    {
                        "name": c.name,
                        "points_count": info.points_count or 0,
                        "sources": self._sources.get(c.name, []),
                    }
                )
            return out
        except Exception:
            return []

    async def delete_collection(self, collection_name: str):
        self.qdrant.delete_collection(collection_name)
        self._sources.pop(collection_name, None)
